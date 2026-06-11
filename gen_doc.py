"""生成技术文档 Word 文件"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ─── 封面 ───
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n\n')
run = p.add_run('校园留言墙\n技术文档')
run.font.size = Pt(28)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基于 Flask + PostgreSQL 的全栈 Web 应用')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ─── 1. 项目概述 ───
doc.add_heading('一、项目概述', level=1)
doc.add_paragraph(
    '校园留言墙是一个面向在校学生的轻量级社交平台，支持留言发布、互动回复、点赞等功能。'
    '采用单体架构设计，便于快速部署和维护。'
)

doc.add_heading('核心功能', level=2)
for f in [
    '用户注册/登录（学号认证）',
    '留言发布（标题、正文、图片三选一即可）',
    '嵌套回复（可以对回复进行回复）',
    '点赞互动（AJAX 无刷新）',
    '全文搜索（同时搜索标题和正文）',
    '管理员后台（置顶、删除留言）',
    '头像上传（base64 存数据库，永不丢失）',
    '夜间模式切换',
]:
    doc.add_paragraph(f, style='List Bullet')

# ─── 2. 技术栈 ───
doc.add_heading('二、技术栈', level=1)
table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['层级', '技术', '说明']):
    table.rows[0].cells[i].text = h
for r, row in enumerate([
    ['前端', 'HTML / CSS / JavaScript', '原生，无框架依赖'],
    ['后端框架', 'Python Flask 3.0', '路由与业务逻辑'],
    ['ORM', 'Flask-SQLAlchemy 3.1.1', 'Python 操作数据库'],
    ['数据库', 'PostgreSQL 13 + pg8000', '数据持久化存储'],
    ['部署', 'Railway + Gunicorn', '云托管与 WSGI 服务器'],
    ['版本控制', 'Git + GitHub', '代码托管'],
]):
    for c, v in enumerate(row):
        table.rows[r+1].cells[c].text = v

# ─── 3. 项目结构 ───
doc.add_heading('三、项目结构', level=1)
doc.add_paragraph(
    'message_board/\n'
    '├── app.py            主程序（路由 + 业务逻辑）\n'
    '├── models.py         数据库模型定义\n'
    '├── config.py         应用配置\n'
    '├── requirements.txt  Python 依赖\n'
    '├── Procfile          Railway 部署配置\n'
    '├── static/\n'
    '│   ├── style.css     全局样式\n'
    '│   └── compress.js   前端图片压缩\n'
    '├── templates/        HTML 模板（11 个）\n'
    '│   ├── base.html     母版页\n'
    '│   ├── index.html    首页\n'
    '│   ├── publish.html  发布页\n'
    '│   ├── message.html  详情+回复\n'
    '│   ├── login.html    登录\n'
    '│   ├── register.html 注册（验证码）\n'
    '│   ├── profile.html  设置+头像\n'
    '│   └── ...\n'
    '└── uploads/          上传目录'
)

# ─── 4. 数据库设计 ───
doc.add_heading('四、数据库设计', level=1)
doc.add_paragraph('共 4 张表，通过外键关联：')

tables_desc = [
    ('users（用户表）', [
        'id: 主键，自增',
        'student_id: 学号（唯一）',
        'username: 用户名（登录用）',
        'password_hash: 密码哈希（scrypt）',
        'nickname: 显示昵称',
        'avatar_url: 头像（支持 base64）',
        'is_admin: 是否管理员',
    ]),
    ('messages（留言表）', [
        'id: 主键',
        'user_id: 外键 -> users.id',
        'title: 标题',
        'content: 正文',
        'image_url: 图片（base64）',
        'is_pinned: 是否置顶',
    ]),
    ('replies（回复表）', [
        'id: 主键',
        'message_id: 外键 -> messages.id',
        'user_id: 外键 -> users.id',
        'parent_id: 外键 -> replies.id（嵌套回复）',
        'content: 回复内容',
    ]),
    ('likes（点赞表）', [
        'id: 主键',
        'user_id + message_id: 联合唯一约束',
        '一人只能点赞一次，再点取消',
    ]),
]
for title, fields in tables_desc:
    doc.add_heading(title, level=2)
    for f in fields:
        doc.add_paragraph(f, style='List Bullet')

# ─── 5. 核心功能 ───
doc.add_heading('五、核心功能实现', level=1)

doc.add_heading('5.1 图片 base64 存储', level=2)
doc.add_paragraph(
    '上传的图片在前端压缩后，后端读取文件二进制数据，转为 base64 编码，'
    '直接存入数据库的 TEXT 字段。不依赖服务器文件系统，部署重启不会丢失。'
)
doc.add_paragraph('前端压缩：compressImage(file, 1200, 0.85) 将图片缩到 1200px 宽')

doc.add_heading('5.2 嵌套回复', level=2)
doc.add_paragraph(
    'Reply 模型包含 parent_id 字段，自引用外键关联到自身。'
    '顶级回复（parent_id IS NULL）显示在主列表，子回复缩进展示。'
)

doc.add_heading('5.3 验证码', level=2)
doc.add_paragraph(
    '基于 Session 的数学验证码：随机生成两个数字，求和后存入 session。'
    '仅在 GET 请求时生成，POST 提交时使用 session 中的值校验，'
    '确保验证码不会在提交时刷新导致误判。'
)

doc.add_heading('5.4 发布防抖', level=2)
doc.add_paragraph(
    '前端 JavaScript 实现：点击发布后按钮立即禁用，显示 5 秒倒计时，'
    '期间所有点击无效。倒计时结束后恢复按钮。'
)

doc.add_heading('5.5 点赞（AJAX）', level=2)
doc.add_paragraph(
    '使用 fetch API 发送 POST 请求，服务器返回 JSON { liked, count }，'
    '前端根据结果更新爱心图标和计数，无需刷新页面。'
)

doc.add_heading('5.6 环境感知配置', level=2)
doc.add_paragraph(
    'config.py 自动识别运行环境：部署在 Railway 时读取系统环境变量'
    ' DATABASE_URL 连接远程 PostgreSQL；本地运行时自动回退到 localhost。'
)

# ─── 6. 部署指南 ───
doc.add_heading('六、部署指南', level=1)

doc.add_heading('本地开发', level=2)
doc.add_paragraph('1. pip install -r requirements.txt')
doc.add_paragraph('2. 确保本地 PostgreSQL 运行，创建数据库 message_board')
doc.add_paragraph('3. python app.py')
doc.add_paragraph('4. 浏览器打开 http://localhost:5000')

doc.add_heading('Railway 部署', level=2)
doc.add_paragraph('1. 推送代码到 GitHub 仓库')
doc.add_paragraph('2. 登录 Railway，创建新项目，选择 Deploy from GitHub')
doc.add_paragraph('3. 添加 PostgreSQL 数据库服务（自动注入环境变量）')
doc.add_paragraph('4. 部署完成后自动生成 URL，即可访问')

# ─── 7. API 接口 ───
doc.add_heading('七、API 接口一览', level=1)
api_list = [
    ['GET', '/', '首页留言列表', '否'],
    ['GET', '/publish', '发布页面', '否'],
    ['POST', '/publish', '提交留言', '是'],
    ['GET', '/message/<id>', '留言详情', '否'],
    ['POST', '/message/<id>', '提交回复', '是'],
    ['POST', '/message/<id>/delete', '删除留言', '是'],
    ['POST', '/like/<id>', '点赞/取消', '是'],
    ['GET', '/search?q=', '搜索留言/用户', '否'],
    ['POST', '/register', '用户注册', '否'],
    ['POST', '/login', '用户登录', '否'],
    ['GET', '/logout', '退出登录', '是'],
    ['GET', '/profile', '个人信息页', '是'],
    ['POST', '/profile', '更新信息/头像', '是'],
    ['GET', '/admin', '管理后台', '管理员'],
    ['POST', '/admin/pin/<id>', '置顶/取消', '管理员'],
]
table3 = doc.add_table(rows=len(api_list)+1, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['方法', '路由', '说明', '权限']):
    table3.rows[0].cells[i].text = h
for r, row in enumerate(api_list):
    for c, v in enumerate(row):
        table3.rows[r+1].cells[c].text = v

# ─── 8. 常见问题 ───
doc.add_heading('八、常见问题', level=1)

qa = [
    ('图片上传后消失？',
     '改用 base64 存储后已解决。图片直接存在 PostgreSQL 数据库的 TEXT 字段中，'
     '不依赖服务器磁盘，任何部署重启或重新部署都不会丢失。'),
    ('Railway 部署后连不上数据库？',
     'Railway 提供的是 postgresql:// 格式的 URL，需要在 config.py 中转为 '
     'postgresql+pg8000:// 格式，代码已自动处理。'),
    ('管理员账号是什么？',
     '应用首次启动时自动创建：用户名 admin，密码 admin123。可在管理后台操作。'),
    ('如何迁移数据？',
     '项目提供了一键迁移 Python 脚本，支持从 Railway PostgreSQL 导出全部数据'
     '并导入到本地 PostgreSQL，包括用户、留言、回复、点赞。'),
    ('可以在手机上用吗？',
     'CSS 已适配移动端响应式布局，手机浏览器直接打开网址即可使用。'),
]
for q, a in qa:
    doc.add_heading(q, level=2)
    doc.add_paragraph(a)

# ─── 保存 ───
path = os.path.join(os.path.expanduser('~'), 'Desktop', '校园留言墙_技术文档.docx')
doc.save(path)
print(f'✅ 已保存到桌面: {path}')
